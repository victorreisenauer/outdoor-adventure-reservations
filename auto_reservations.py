"""this script will access take the safaris in the input folder, extract the data for each day of 
each safari and return a document containing pre written emails for each day of the safari for each 
safari"""


# --------------------------------------------script begin ---------------------------------------
import os
import xlrd
from docx import Document
from datetime import datetime, timedelta
os.chdir(r'C:\Users\Victor\automate_reservations\safaris') 
# directs us to the respective folder where the data is stored

list_of_safaris = os.listdir(r'C:\Users\Victor\automate_reservations\safaris')
test_list_of_safaris = list_of_safaris[:2]									# !!!!!! only for experimenting purposes
#list of safaris available for input - used to create safari objects

def get_cell_content(safari_doc_name, row, column):
	workbook = xlrd.open_workbook(safari_doc_name)
	sheet = workbook.sheet_by_index(0)
	try:
		cell_cont = sheet.cell_value(row-1, column-1)
		return cell_cont					# returns the cell's content in a the natural format
	except IndexError:
		return None

# ---------------------------------------------classes begin----------------------------------------

#class SAFARI gets the year, name and duration data about the safari objects and is a queue for the lodge nodes
class Safari:
	def __init__(self, safari_doc_name):
		self.safari_doc_name = safari_doc_name
		self.year = None
		self.name = None
		self.duration = 0
		self.head = None
		self.tail = None
		self.size = 0 
		self.days_to_stay = 1

	# helper functions
	def get_cell_content(self, row, column):
		return get_cell_content(self.safari_doc_name, row, column)

	def get_size(self):
		return self.size

	def is_empty(self):
		return self.size == 0 
	# end of helper functions

	def peek(self):
		if self.size > 0:
			return self.head
		return "No lodge here"

	def enqueue(self):
		skip_forward_helper = 0
		for day in range(1, self.get_safari_data()["safari_duration"] + 1):
			if skip_forward_helper > 0:
				skip_forward_helper -=1
				continue
			days_to_stay = 1
			lodge_node = LodgeNode(day, self.safari_doc_name)
			switch = True
			while switch:
				if lodge_node.get_lodge_name() == lodge_node.get_next_lodges_name(days_to_stay):
					days_to_stay += 1
					skip_forward_helper = days_to_stay-1
				else:
					switch = False

			if self.is_empty():
				self.head = lodge_node
				self.tail = lodge_node 
			else:
				self.tail.set_next_lodge(lodge_node)
				self.tail = lodge_node
			self.size += 1
			#print(self.tail.get_lodge_name())
			#print(days_to_stay)
			lodge_node.set_number_of_days(days_to_stay)

			
			#print(self.days_to_stay)


			#print("first lodge: {0}		| last lodge: {1}".format(self.head, self.tail))

	def dequeue(self):
		if self.is_empty():
			return "No more lodges"

		if self.head.day == self.get_safari_data()["safari_duration"]:
			return "Guests are back in Windhoek or Kathima"

		lodge_to_remove = self.head
		#print("Removing {0}".format(self.head))
		if self.size == 1:
			self.head = None
			self.tail = None
		else:
			self.head = lodge_to_remove.get_next_lodge()
			#print("The new head is: {0}".format(self.head))
		self.size -= 1
		return lodge_to_remove.get_value()

	def get_safari_data(self):					#used to access the data of that safari object
		lst = xlrd.xldate_as_tuple(self.get_cell_content(19, 3), 0)
		date = datetime(year=int(lst[0]), month=int(lst[1]), day=int(lst[2])).date()
		self.year = date.strftime("%Y")
		self.name = self.get_cell_content(3, 1)
		self.duration = 0
		switch = True	
		while switch:
			if self.get_cell_content(19 + self.duration, 1) is None:
				switch = False
			self.duration += 1
		safari_data = {"safari_year": self.year, "safari_name": self.name, "safari_duration": self.duration-1}	
		return safari_data	


#----------------------------------------------------------------------------------------

"""
test_safari = Safari(test_list_of_safaris[0])
print(test_safari.get_safari_data())
"""

#----------------------------------------------------------------------------------------
# LodgeNode will create a node for each day's lodge that we can queue later on
class LodgeNode:
	def __init__(self, day, safari_doc_name, next_lodge=None):
		#self.day = [i for i in range(1, duration+1)]
		self.day = day
		self.num_of_days = 1
		self.lodge_name = None
		self.contact_person = None
		self.safari_doc_name = safari_doc_name
		self.next_lodge = next_lodge
		self.value = {}

	# helper functions start
	def __repr__(self):
		self.lodge_name = self.get_lodge_name()
		return self.lodge_name

	def get_cell_content(self, row, column, special=False):
		if not special:
			return get_cell_content(self.safari_doc_name, row, column)
		os.chdir(r'C:\Users\Victor\automate_reservations\lodge_infos')
		value = get_cell_content("data_auto_reservations.xlsx", row, column)
		os.chdir(r'C:\Users\Victor\automate_reservations\safaris')
		return value
	# helper functions end


	def set_next_lodge(self, next_lodge):
		self.next_lodge = next_lodge

	def get_next_lodge(self):
		return self.next_lodge	
		
	def get_lodge_name(self):							#get the name of the lodge | add is used to get next lodge
		return self.get_cell_content(18 + self.day, 5)

	def get_next_lodges_name(self, add=1):
		return self.get_cell_content(18 + self.day + add, 5)

	def get_acc_type(self):
		acc_type = self.get_cell_content(18 + self.day, 8)
		return acc_type

	def get_value(self):
		self.value["lodge_name"] = self.get_lodge_name()
		self.value["contact_person"] = self.get_contact_person()
		self.value["day_in"] = self.get_day_in()
		self.value["day_out"] = self.get_day_out()
		self.value["number_of_days"] = self.get_number_of_days()
		self.value["acc_type"] = self.get_acc_type()
		return self.value

	def get_contact_person(self):						#get the contact person's name of that lodge
		j = 0
		while True:
			j += 1
			iterate_lodges = self.get_cell_content(4+j, 3, True)
			if self.get_lodge_name() == iterate_lodges:
				returned_value = self.get_cell_content(4+j, 9, True)
				break
			elif iterate_lodges is None:
				returned_value = ("lodge not found in the list of lodges")
				break
		return returned_value

	def get_day_in(self):								#get the arrival day at the lodge
		#print(self.day)
		try:
			lst = xlrd.xldate_as_tuple(self.get_cell_content(18 + self.day, 3), 0)
			date = datetime(year=int(lst[0]), month=int(lst[1]), day=int(lst[2])).date()
			year = date.strftime("%Y")
			day = date.strftime("%d")
			month = date.strftime("%b")
			return "{0}/{1}/{2}".format(day, month, year)
		except TypeError:
			return "Did you input dates for the safari?"

	def get_day_out(self):								#get the departure day at the lodge
		num_of_days = self.get_number_of_days()
		lst = xlrd.xldate_as_tuple(self.get_cell_content(18 + self.day + num_of_days, 3), 0)
		date = datetime(year=int(lst[0]), month=int(lst[1]), day=int(lst[2])).date()
		year = date.strftime("%Y")
		day = date.strftime("%d")
		month = date.strftime("%b")
		return "{0}/{1}/{2}".format(day, month, year)

	def set_number_of_days(self, num_of_days):
		self.num_of_days = num_of_days

	def get_number_of_days(self):						#get the number_of_days we stayed at the lodge
		return self.num_of_days
#--------------------------------------------------------------------------
# testing LodgeNode functionality:

#test_lodge = LodgeNode(1, test_list_of_safaris[0])
#print(test_lodge.get_day_in())
#print(test_lodge.get_day_out())



#print(LodgeNode(1, test_list_of_safaris[1]).get_value())
#---------------------------------------------------------------------------


class Template:
	def __init__(self, lodge_values, safari):
		self.contact_person = lodge_values["contact_person"]
		self.year = safari.year
		self.tour_name = safari.name
		self.lodge_name = lodge_values["lodge_name"]
		self.date_in = lodge_values["day_in"]
		self.date_out = lodge_values["day_out"]
		self.num_nights = lodge_values["number_of_days"]
		self.acc_type = lodge_values["acc_type"]



	def template_insert(self):
		return """
Dear {0},

Please make a reservation for {1}.

Reference name: {2}.

date in: {3}
date out: {4}
number of nights: {5}

{6} | 7 {7} please.
min. 6 - max. 10 persons
+ 1 guideroom

Please confirm asap!

Best regards,
Alina
		""".format(self.contact_person, self.year, self.tour_name, self.date_in, \
			self.date_out, self.num_nights, self.lodge_name, self.acc_type)


#----------------------------------------------------------------------------------------

"""
for i in range(len(test_list_of_safaris)):
	safari = Safari(test_list_of_safaris[i])
	safari.enqueue()
	for day in range(1, safari.size): # deliberately not "+1" to not print out the last station (windhoek or kathima)
		lodge_values = safari.dequeue()
		email = Template(lodge_values)
		print(email.template_insert())
		document = Document()
		message = document.add_paragraph(email.template_insert())
		document.save("emails")

		
		#print("")
		#print("")
		#print("")

"""
def create_documents():
	for i in range(len(test_list_of_safaris)):
		safari = Safari(test_list_of_safaris[i])
		safari.enqueue()
		name = safari.get_safari_data()["safari_name"].split(" ")
		safari_name = name[0] + ".docx"
		os.chdir(r'C:\Users\Victor\automate_reservations\created_reservation_emails')
		document = Document()
		document.save(safari_name)
		os.chdir(r'C:\Users\Victor\automate_reservations\safaris') 
		for day in range(1, safari.size): # deliberately not "+1" to not print out the last station (windhoek or kathima)
			lodge_values = safari.dequeue()
			email = Template(lodge_values, safari)
			print(email.template_insert())
			os.chdir(r'C:\Users\Victor\automate_reservations\created_reservation_emails')
			document = Document(safari_name)
			document.add_paragraph(email.template_insert())
			document.add_paragraph("-----------------------------------------------------------------------------")
			document.save(safari_name)
			os.chdir(r'C:\Users\Victor\automate_reservations\safaris') 

create_documents()

